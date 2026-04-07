# -*- coding: utf-8 -*-
"""
发布日志生成脚本
用法：python gen_release_log.py <Excel文件路径> <Sheet名称> [输出文件名]
依赖：pip install python-docx pandas openpyxl
"""

import sys
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

# ─── 项目配置 ────────────────────────────────────────────────────
PROJECT_ORDER = [
    '个人产品',
    '寄件联盟',
    'API开放平台',
    '商家寄件',
    '企业快递管理SaaS',
    '电商快递管家SaaS',
    '收件端',
    '业务平台',
]
PROJECT_KEYS = {
    '个人产品':           '个人产品(#5)',
    '寄件联盟':           '寄件联盟(#23)',
    'API开放平台':        'API开放平台(#12)',
    '商家寄件':           '商家寄件(#29)',
    '企业快递管理SaaS':   '企业快递管理SaaS(#19)',
    '电商快递管家SaaS':   '电商快递管家SaaS(#1)',
    '收件端':             '收件端(#4)',
    '业务平台':           '业务平台(#37)',
}
COL_WIDTHS  = [544, 2417, 2306, 3029]
BLUE_COLOR  = '0082FF'
OUTPUT_DIR  = 'D:/office/项目数据汇总/发布日志/2026/'
TEMPLATE    = 'D:/office/项目数据汇总/发布日志/2026/2026年3月31日发布日志.docx'

# ─── 辅助函数 ────────────────────────────────────────────────────
def set_run_font(run, size_pt=12, bold=None, color_hex=None):
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'),    '微软雅黑')
    rFonts.set(qn('w:hAnsi'),    '微软雅黑')
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rFonts.set(qn('w:cs'),       '微软雅黑')
    sz_val = str(int(size_pt * 2))
    for tag in ('w:sz', 'w:szCs'):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn('w:val'), sz_val)
    if bold is not None:
        for btag in ('w:b', 'w:bCs'):
            el = rpr.find(qn(btag))
            if bold:
                if el is None:
                    el = OxmlElement(btag)
                    rpr.append(el)
            else:
                if el is not None:
                    rpr.remove(el)
    if color_hex:
        color_el = rpr.find(qn('w:color'))
        if color_el is None:
            color_el = OxmlElement('w:color')
            rpr.append(color_el)
        color_el.set(qn('w:val'), color_hex)


def set_para_line_spacing(para, spacing):
    pPr = para._p.get_or_add_pPr()
    lSpacing = pPr.find(qn('w:spacing'))
    if lSpacing is None:
        lSpacing = OxmlElement('w:spacing')
        pPr.append(lSpacing)
    lSpacing.set(qn('w:line'),     str(int(spacing * 240)))
    lSpacing.set(qn('w:lineRule'), 'auto')


def set_cell_para_line_spacing(cell, line_twips=400):
    for p in cell.paragraphs:
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p._p.insert(0, pPr)
        spacing_el = pPr.find(qn('w:spacing'))
        if spacing_el is None:
            spacing_el = OxmlElement('w:spacing')
            pPr.append(spacing_el)
        spacing_el.set(qn('w:line'),     str(line_twips))
        spacing_el.set(qn('w:lineRule'), 'exact')


def set_cell_valign(cell, val='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        tcPr.append(vAlign)
    vAlign.set(qn('w:val'), val)


def set_cell_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'),    str(width_twips))
    tcW.set(qn('w:type'), 'dxa')


def set_row_height(row, val_twips=360):
    trPr = row._tr.get_or_add_trPr()
    trH = trPr.find(qn('w:trHeight'))
    if trH is None:
        trH = OxmlElement('w:trHeight')
        trPr.append(trH)
    trH.set(qn('w:val'), str(val_twips))


def add_border_element(parent, tag, val='single', sz='4', space='0', color='000000'):
    el = OxmlElement(tag)
    el.set(qn('w:val'),   val)
    el.set(qn('w:sz'),    sz)
    el.set(qn('w:space'), space)
    el.set(qn('w:color'), color)
    parent.append(el)
    return el


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV'):
        add_border_element(tblBorders, side)
    tblPr.append(tblBorders)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('w:top', 'w:left', 'w:bottom', 'w:right'):
        add_border_element(tcBorders, side)
    tcPr.append(tcBorders)


# ─── 主函数 ───────────────────────────────────────────────────────
def generate_release_log(excel_path, sheet_name, release_date_str='2026年3月31日', output_filename=None):
    # 读取数据
    df = pd.read_excel(excel_path, sheet_name=sheet_name)
    df.columns = df.columns.str.strip()
    released = df[df['需求状态'] == '已发布'].copy()
    released = released.drop_duplicates(subset=['禅道编号'])

    # 按模块分组
    ordered_projects = []
    total = 0
    for pname in PROJECT_ORDER:
        key = PROJECT_KEYS.get(pname, pname)
        group = released[released['关联项目'] == key]
        if len(group) == 0:
            continue
        items = []
        for seq, (_, row) in enumerate(group.iterrows(), start=1):
            items.append({
                'seq':      seq,
                'name':     str(row.get('需求名称', '')).strip(),
                'category': str(row.get('需求分类', '')).strip(),
                'value':    str(row.get('需求意义与预期价值说明', '')).strip(),
            })
        total += len(items)
        ordered_projects.append({'name': pname, 'count': len(items), 'items': items})

    if not ordered_projects:
        print('未找到任何已发布需求，请检查数据和筛选条件。')
        return

    # 构建文档
    doc = Document(TEMPLATE)
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'sectPr':
            body.remove(child)

    # 标题段
    p_title = doc.add_paragraph(style='Heading 3')
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(f'{release_date_str}发布日志')
    set_run_font(run, size_pt=12, bold=None)

    # 总结段
    proj_list_str = '、'.join([p['name'] for p in ordered_projects])
    summary_text = (
        f'{release_date_str}迭代版本总共发布了{total}个需求，'
        f'包括新功能及产品优化。本次发布内容包括{proj_list_str}等。具体发布内容如下：'
    )
    p_summary = doc.add_paragraph(style='Normal')
    set_para_line_spacing(p_summary, 2.0)
    run = p_summary.add_run(summary_text)
    set_run_font(run, size_pt=12, bold=False)

    # 各模块
    CN_NUM = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    for idx, proj in enumerate(ordered_projects):
        # 空行
        p_blank = doc.add_paragraph(style='Normal')
        set_para_line_spacing(p_blank, 1.5)
        run_blank = p_blank.add_run(' ')
        set_run_font(run_blank, size_pt=12)

        # 模块标题
        cn = CN_NUM[idx] if idx < len(CN_NUM) else str(idx + 1)
        title_text = f'（{cn}）{proj["name"]}（{proj["count"]}个）'
        p_mod = doc.add_paragraph(style='Normal')
        set_para_line_spacing(p_mod, 1.5)
        run_mod = p_mod.add_run(title_text)
        set_run_font(run_mod, size_pt=12, bold=True, color_hex=BLUE_COLOR)

        # 表格
        table = doc.add_table(rows=1 + len(proj['items']), cols=4)
        table.style = doc.styles['Normal Table']
        set_table_borders(table)

        # 表头
        hrow = table.rows[0]
        set_row_height(hrow, 360)
        header_labels = ['序号', '需求名称', '需求分类', '需求意义与预期价值']
        for ci, label in enumerate(header_labels):
            cell = hrow.cells[ci]
            set_cell_valign(cell, 'center')
            set_cell_width(cell, COL_WIDTHS[ci])
            set_cell_borders(cell)
            set_cell_para_line_spacing(cell, 400)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(label)
            set_run_font(run, size_pt=12, bold=True)

        # 数据行
        for ri, item in enumerate(proj['items']):
            drow = table.rows[ri + 1]
            set_row_height(drow, 360)
            values = [str(item['seq']), item['name'], item['category'], item['value']]
            for ci, val in enumerate(values):
                cell = drow.cells[ci]
                set_cell_valign(cell, 'center')
                set_cell_width(cell, COL_WIDTHS[ci])
                set_cell_borders(cell)
                set_cell_para_line_spacing(cell, 400)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val)
                set_run_font(run, size_pt=12, bold=False, color_hex='000000')

    # 保存
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    if output_filename is None:
        base = os.path.splitext(os.path.basename(excel_path))[0]
        output_filename = f'{base}(整理版-已发布).docx'
    out_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(out_path)
    print(f'完成！共 {total} 个需求，保存至: {out_path}')
    return out_path


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('用法: python gen_release_log.py <Excel文件> <Sheet名称> [发布日期如"2026年3月31日"] [输出文件名]')
        sys.exit(1)
    excel_path  = sys.argv[1]
    sheet_name  = sys.argv[2]
    release_date = sys.argv[3] if len(sys.argv) > 3 else '2026年3月31日'
    out_file    = sys.argv[4] if len(sys.argv) > 4 else None
    generate_release_log(excel_path, sheet_name, release_date, out_file)
